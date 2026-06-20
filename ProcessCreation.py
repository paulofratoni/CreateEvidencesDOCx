import os
import glob
from openpyxl import load_workbook
from docx import Document

# --- INITIAL CONFIG ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FOLDER = os.path.join(os.path.expanduser("~"), "Downloads", "Evidences", "Input")
OUTPUT_FOLDER = os.path.join(os.path.expanduser("~"), "Downloads", "Evidences", "Output")
INPUT_PATTERN = "BDD*.xlsx"  # Pattern for the XLSX file(s)
MODEL_NAME = "modelo.docx"

# Create the output folder if it doesn't exist
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# Look for the spreadsheet file using the provided input folder
excel_files = glob.glob(os.path.join(INPUT_FOLDER, INPUT_PATTERN))
if not excel_files:
    raise FileNotFoundError(
        f"Não foi encontrado nenhum arquivo correspondente a {INPUT_PATTERN} em {INPUT_FOLDER}"
    )

excel_path = excel_files[0]

# Look for a template in the input folder; if not found, create a simple one
model_candidates = [
    os.path.join(INPUT_FOLDER, MODEL_NAME),
    *glob.glob(os.path.join(INPUT_FOLDER, "*.docx")),
    os.path.join(BASE_DIR, MODEL_NAME),
]
PATH_MODEL = next((path for path in model_candidates if path and os.path.exists(path)), None)

if PATH_MODEL is None:
    PATH_MODEL = os.path.join(INPUT_FOLDER, MODEL_NAME)
    template_doc = Document()
    template_doc.add_heading("Template padrão", level=1)
    template_doc.add_paragraph("Modelo para geração dos documentos.")
    template_doc.save(PATH_MODEL)

# --- LEITURA DOS DADOS ---
# Lendo a planilha com os cabeçalhos na segunda linha do arquivo
# (o layout do Excel fornecido usa essa configuração)
wb = load_workbook(excel_path, data_only=True)
sheet = wb.active

header_row = 2
header_map = {}
for col_idx in range(1, sheet.max_column + 1):
    header_value = sheet.cell(row=header_row, column=col_idx).value
    if header_value is not None:
        header_map[str(header_value).strip()] = col_idx

print("Iniciando a geração dos arquivos Word...\n")

# --- PROCESSAMENTO POR LINHA ---
for row_idx in range(header_row + 1, sheet.max_row + 1):
    # Ignore blank or untitled lines.
    row_values = [sheet.cell(row=row_idx, column=col_idx).value for col_idx in range(1, sheet.max_column + 1)]
    if all(value is None or str(value).strip() == "" for value in row_values):
        continue

    def get_value(column_name):
        col_idx = header_map.get(column_name)
        if col_idx is None:
            return ""
        value = sheet.cell(row=row_idx, column=col_idx).value
        return "" if value is None else str(value)

    title_value = get_value("Title")
    if title_value.strip() == "":
        continue

    file_title = title_value.strip()
    given_line = get_value("Given")
    when_line = get_value("When")
    then_line = get_value("Then")
    
    # 1. Open a new copy of the model document
    doc = Document(PATH_MODEL)
    
    # 2. Insert a page break at the end of the model to go to the 2nd page
    doc.add_page_break()
    
    # 3. Add the contents to the seconde page.
    # Add the title as a header in the document
    doc.add_heading(f"Cenário: {file_title}", level=1)
    
    # Add the Given, When, Then blocks
    #doc.add_heading("Given (Dado que):", level=2)
    doc.add_paragraph(given_line)
    
    #doc.add_heading("When (Quando):", level=2)
    doc.add_paragraph(when_line)
    
    #doc.add_heading("Then (Então):", level=2)
    doc.add_paragraph(then_line)
    
    # 4. Save the new file with the name from the 'Title' column
    # Preserve the original text as much as possible, only removing characters that cannot be used in file names
    invalid_chars = '<>:"/\\|?*\x00'
    clean_name = "".join(ch for ch in file_title if ch not in invalid_chars)
    clean_name = clean_name.replace("\r", " ").replace("\n", " ")

    # Fallback only if the title becomes empty after sanitization
    if not clean_name:
        clean_name = f"Cenario_{row_idx}"

    save_path = os.path.join(OUTPUT_FOLDER, f"{clean_name}.docx")
    
    doc.save(save_path)
    print(f"[OK] Arquivo gerado: {save_path}")

print("\nProcess completed successfully!")